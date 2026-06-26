#!/usr/bin/env python3
"""Build the V-LiSEMOD manuscript Markdown and DOCX drafts.

This script is intentionally self-contained so collaborators can regenerate the
draft after editing source text or replacing figure assets. It uses only
repository-local manuscript notes plus python-docx/Pillow.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except Exception as exc:  # pragma: no cover - reported in main
    Document = None
    IMPORT_ERROR = exc
else:
    IMPORT_ERROR = None

try:
    from PIL import Image
except Exception:
    Image = None


ROOT = Path(__file__).resolve().parents[1]
DRAFT_DIR = ROOT / "docs" / "manuscript" / "drafts"
FIG_DIR = ROOT / "docs" / "manuscript" / "figures"
MD_PATH = DRAFT_DIR / "V-LiSEMOD_manuscript_draft.md"
DOCX_PATH = DRAFT_DIR / "V-LiSEMOD_manuscript_draft.docx"
NOTES_PATH = DRAFT_DIR / "MANUSCRIPT_DRAFT_NOTES.md"


SOURCE_FILES = [
    "README.md",
    "docs/APP_GUIDE.md",
    "docs/DATABASE.md",
    "docs/DEVELOPER_NOTES.md",
    "docs/PROTACABILITY.md",
    "docs/MANUSCRIPT_OUTLINE.md",
    "docs/manuscript/MANUSCRIPT_PLAN.md",
    "docs/manuscript/CLAIMS_AND_LIMITATIONS_MATRIX.md",
    "docs/manuscript/FIGURE_AND_TABLE_PLAN.md",
    "docs/manuscript/VALIDATION_AND_REPRODUCIBILITY_PLAN.md",
    "docs/manuscript/figures/FIGURE_INDEX.md",
    "docs/manuscript/figures/FIGURE_CAPTIONS.md",
    "docs/manuscript/figures/FIGURE_STORYBOARD.md",
    "docs/manuscript/figures/SCREENSHOT_CAPTURE_NOTES.md",
    "docs/manuscript/figures/IMAGE_CLEANUP_NOTES.md",
    "docs/manuscript/figures/TOOL_URL_CHECKS.md",
    "docs/manuscript/figures/api_evidence/README.md",
]


@dataclass(frozen=True)
class Figure:
    number: int
    title: str
    files: tuple[str, ...]
    caption: str
    placeholder: str | None = None


FIGURES = [
    Figure(
        1,
        "Connected companion-tool ecosystem for structure-guided ligand and degrader-oriented workflows",
        ("annotated/fig01_ecosystem_overview_draft.png",),
        "Draft figure placeholder. The current ecosystem montage illustrates Warhead Hunter, V-LiSEMOD, PROTAC Builder, E3 Ligandalyzer, and PyMACS as companion resources for adjacent stages of structure-guided analysis. It supports ecosystem framing and downstream workflow continuity, but it does not establish full technical integration, synchronized execution across tools, or a single shared backend.",
    ),
    Figure(
        2,
        "AutoDock-Vina PrepServer web workflow",
        (),
        "Draft figure placeholder. This slot is reserved for the docking-package preparation tool described in the planning materials. The present workspace is V-LiSEMOD rather than an AutoDock-Vina PrepServer repository, so no screenshot is inserted in this draft.",
        placeholder="Image pending: regenerate this panel once the correct local AutoDock-Vina PrepServer application or repository is available.",
    ),
    Figure(
        3,
        "Warhead Hunter public workflow and results-library context",
        (
            "cropped/fig03_warhead_hunter_panel.png",
            "cropped/fig03_warhead_exposure_result_panel.png",
        ),
        "Draft figure placeholder. Warhead Hunter is shown through a public workflow-launch panel and a public results-library crop. These screenshots demonstrate inspectable, exposure-oriented workflow surfaces, but they do not prove validated attachment-site selection, a guaranteed linker position, or automated PROTAC design success.",
    ),
    Figure(
        4,
        "V-LiSEMOD viral ligand evidence review and PROTACability-style triage",
        (
            "cropped/fig04_vlisemod_home_panel.png",
            "cropped/fig04_vlisemod_protein_query_panel.png",
            "cropped/fig04_vlisemod_ligand_indexer_panel.png",
            "cropped/fig04_vlisemod_compare_ligands_panel.png",
            "cropped/fig04_vlisemod_protacability_filters_results.png",
        ),
        "Draft figure placeholder. V-LiSEMOD is represented by Structure Explorer, Protein Query, Ligand Indexer, Ligand Comparison, and PROTACability-style triage panels. The figure supports a hypothesis-generation and structure-guided review claim. PROTACability outputs are structural-priority heuristics and should not be interpreted as experimentally validated degradation predictions or guaranteed target degradability.",
    ),
    Figure(
        5,
        "PROTAC Builder as a downstream continuation workspace",
        (
            "cropped/fig05_protacbuilder_home_panel.png",
            "cropped/fig05_protacbuilder_design_workspace.png",
        ),
        "Draft figure placeholder. These panels position PROTAC Builder as a companion continuation tool for downstream warhead, linker, and recruiter assembly after upstream structural triage. The figure does not show experimental validation, guaranteed design success, or a fully unified automated pipeline.",
    ),
    Figure(
        6,
        "E3 Ligandalyzer and PyMACS as companion context resources",
        ("cropped/fig06_e3_pymacs_companion_panels.png",),
        "Draft figure placeholder. E3 Ligandalyzer provides recruiter-centered companion context, while PyMACS is represented as a companion computational repository. The figure should not be read as evidence that either resource is embedded directly into V-LiSEMOD or exposed through the same public web service.",
    ),
    Figure(
        7,
        "Public API and reproducibility evidence surfaces",
        ("raw/fig07_warhead_api_docs_full.png",),
        "Draft figure placeholder. The screenshot and captured endpoint files document public API/documentation surfaces and lightweight health or manifest evidence at the time of capture. They should not be interpreted as production-scale API guarantees, long-term endpoint stability, or service-level commitments.",
    ),
]


TABLES = {
    "Platform capabilities": [
        ["Module", "Primary user question", "Manuscript-safe role"],
        ["Structure Explorer", "What does this ligand-bound viral structure show?", "Structure-specific review, ligand imagery, PyMOL-oriented export, exposed-atom and functional-group context when available."],
        ["Protein Query", "Which structures exist for a target or virus?", "Target-centric filtering, PDB selection, export-oriented data assembly, and PROTACability handoff."],
        ["Ligand Indexer", "Where does this ligand appear?", "Ligand-first lookup across mapped viral structure contexts with interaction chart generation."],
        ["Ligand Comparison", "How does a ligand behave across structures?", "Cross-structure comparison of ligand contacts and mapped context, not a validated SAR engine."],
        ["PROTACability Assessment", "Which contexts merit degrader-oriented review?", "Transparent heuristic triage across target, protein, structure, and chain views."],
        ["Companion-tool handoff", "Where can follow-up design exploration continue?", "External continuation into PROTAC Builder and related tools without implying full backend unification."],
        ["Optional assistant route", "Can a deployment provide local assistant context?", "Deployment-dependent surface only when explicitly enabled."],
    ],
    "Data-layer summary": [
        ["Layer", "Representative tables or assets", "Role in manuscript framing"],
        ["Target metadata", "Virus_Proteins", "Curated virus, protein, and PDB-linked target context."],
        ["Ligand identity", "ligand_synonyms, Ligand_Atoms_Smiles, SMILES_MAP_PDB", "Ligand-code, synonym, SMILES, and mapping support for lookup and comparison."],
        ["Structural atoms and pockets", "ligand_atoms, distal_atoms, receptor_binding_pocket", "Atom-level ligand context and binding-pocket evidence for selected structures."],
        ["Interactions", "Arpeggio_Contacts_Data, Ligand_Arp_Diagram", "Descriptive contact summaries and interaction charts."],
        ["Solvent exposure", "RUPLEY_SASA_DATA, solvent_exposed_atoms, ligand_water_distances", "SASA-derived exposed-atom and solvent-context cues."],
        ["Functional groups", "Functional_GROUPED, Functional_Group_Atoms", "Functional-group annotations for ligand-context interpretation."],
        ["PROTACability", "protacability_assessment and related lysine, ligand inventory, warhead linkability, and readiness tables", "Transparent degrader-readiness triage layers rather than experimental degradation evidence."],
        ["Generated assets", "static/charts, static/ligand_images, pml_sessions, output_csvs", "Runtime outputs and exports that require writable local storage."],
    ],
    "PROTACability interpretation guardrails": [
        ["Evidence layer", "Useful interpretation", "What it does not establish"],
        ["Warhead linkability", "Ligand-centered attachment opportunity based on exposed atoms, contact preservation, and functional-group context.", "An optimal linker atom or confirmed medicinal chemistry route."],
        ["Target lysine accessibility", "Target-side exposed lysine cues in the same structural context.", "Ubiquitination, degradation, or productive ternary-complex formation."],
        ["Protein structural priority", "A higher-priority starting point for design review.", "Target degradability or biological tractability."],
        ["Ternary geometry cue", "A simplified geometry-informed hypothesis for follow-up.", "A modeled ternary complex or binding-energy prediction."],
        ["Overall degrader readiness", "A combined structural-priority heuristic for triage.", "Experimentally validated degrader prediction or guaranteed PROTAC success."],
    ],
    "Module-to-user-question map": [
        ["Workflow", "Representative question", "Output to review"],
        ["Structure-specific review", "Which ligand atoms appear exposed enough to discuss as modification sites?", "Ligand image, exposed-atom context, functional-group annotations, and PyMOL-oriented session files."],
        ["Target-centric query", "Which structures and ligands are available for this viral protein?", "Filtered PDB/ligand selections and export-ready linked data."],
        ["Ligand-first lookup", "Where else does this ligand or synonym appear?", "Mapped structures, chains, residue selections, and interaction charts."],
        ["Cross-structure comparison", "Does the same ligand preserve contacts across structures?", "Comparison charts and atom-level interaction burden views."],
        ["Heuristic triage", "Which structures are worth follow-up review for degrader-oriented exploration?", "Filtered PROTACability result tables and evidence-layer detail views."],
    ],
    "Limitations and future work": [
        ["Current limitation", "Manuscript-safe wording", "Future-work direction"],
        ["Co-crystal dependence", "The workflow depends on available ligand-bound structures and prepared annotations.", "Expand curated target and ligand coverage."],
        ["Coverage boundaries", "Database coverage is curated and should not be described as exhaustive.", "Add provenance, release tags, and redacted fixtures."],
        ["Heuristic PROTACability", "PROTACability supports hypothesis generation and design triage.", "Benchmark heuristic layers against curated examples when data are available."],
        ["Ternary-complex simplification", "Geometry cues do not model complete ternary complexes.", "Connect to dedicated modeling and experimental follow-up workflows."],
        ["Optional assistant", "The assistant route is deployment-dependent.", "Clarify enabled/disabled states and fixture tests."],
        ["Generated outputs", "Exports and session files depend on writable local storage.", "Improve packaging and cleanup automation."],
        ["AutoDock-Vina figure slot", "The docking-preparation panel is pending in this draft.", "Recapture once the correct local app is available."],
    ],
    "Reproducibility checklist": [
        ["Item", "Draft status", "Reviewer action"],
        ["Application version or commit", "Placeholder", "Record the exact release, tag, or commit before submission."],
        ["Environment setup", "Partially documented", "Confirm requirements and environment variables for the target deployment."],
        ["Database provisioning", "Documented as local SQLite or RANDY-backed", "Add repository-safe fixture or data-access language."],
        ["Main page validation", "Planned and screenshot-supported", "Re-run route smoke checks for the final version."],
        ["PROTACability imports", "Planned through CSV-to-SQLite importer", "Confirm table presence and dashboard behavior in the final build."],
        ["Figure evidence", "Representative screenshots captured on 2026-06-11", "Recapture any overlay-heavy panels and the pending AutoDock-Vina slot."],
        ["Endpoint evidence", "Lightweight public checks captured", "Avoid service-level or long-term availability claims."],
    ],
}


def md_table(rows: list[list[str]]) -> str:
    header, *body = rows
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join("---" for _ in header) + " |"]
    lines.extend("| " + " | ".join(cell.replace("\n", " ") for cell in row) + " |" for row in body)
    return "\n".join(lines)


def figure_markdown(fig: Figure) -> str:
    lines = [f"### Figure {fig.number}. {fig.title}", ""]
    if fig.placeholder:
        lines.extend([f"**Draft figure placeholder:** {fig.placeholder}", ""])
    for file_name in fig.files:
        path = FIG_DIR / file_name
        if path.exists():
            lines.extend([f"![Figure {fig.number} draft panel](../figures/{file_name})", ""])
        else:
            lines.extend([f"**Missing image:** `docs/manuscript/figures/{file_name}`", ""])
    lines.extend([f"**Caption:** {fig.caption}", ""])
    return "\n".join(lines)


def build_markdown() -> str:
    table_sections = []
    for i, (title, rows) in enumerate(TABLES.items(), start=1):
        table_sections.append(f"### Table {i}. {title}\n\n{md_table(rows)}\n")

    figure_sections = "\n".join(figure_markdown(fig) for fig in FIGURES)

    return f"""# V-LiSEMOD: structure-guided viral ligand analysis with transparent degrader-readiness triage

**Draft manuscript source for collaborator review**

**Working title:** V-LiSEMOD: structure-guided viral ligand analysis with transparent degrader-readiness triage

**Authors:** [AUTHOR NAMES TO ADD]

**Affiliations:** [AFFILIATIONS TO ADD]

**Corresponding author:** [CONTACT TO ADD]

**Running title:** V-LiSEMOD viral ligand analysis platform

**Article type:** Software article, webserver paper, application note, methods article, or resource article. Final article type should be selected after journal targeting.

**Reference status:** Citations in this draft are placeholders for curator review. Do not treat bracketed citation needs as verified bibliography entries.

## Abstract

Viral protein-ligand co-crystal structures contain practical signals for antiviral chemical biology and medicinal chemistry, but design-relevant evidence is often distributed across structural files, interaction tables, solvent accessibility calculations, ligand identifiers, and downstream scripts. V-LiSEMOD is a structure-guided viral protein-ligand exploration platform that brings curated viral target metadata, ligand-centered interaction evidence, solvent-exposed atom review, functional-group annotations, cross-structure ligand comparison, and PROTACability-style heuristic triage into a web-accessible workflow. The platform supports structure-specific exploration, target-centric query, ligand-first lookup, ligand comparison, PyMOL-oriented export, and companion-tool handoff for downstream degrader-oriented design exploration. Its PROTACability layer is framed as transparent hypothesis generation: warhead linkability, target-side lysine accessibility, structural-priority cues, and combined degrader-readiness summaries help users prioritize contexts for review without claiming experimental degradation prediction. Representative screenshots and lightweight endpoint evidence document the current public-facing workflow, while validation notes outline reproducibility checks for local SQLite-backed and optional RANDY-backed deployments. V-LiSEMOD is intended to help researchers move from viral co-crystal context to interpretable design questions, not to replace biochemical validation, cellular degradation assays, ternary-complex modeling, or medicinal chemistry judgment. This draft presents the software architecture, user workflows, data layers, validation plan, limitations, availability placeholders, and manuscript-safe figure package for collaborator review.

## Keywords

viral ligand analysis; structural bioinformatics; solvent exposure; PROTACability; degrader design; ligand comparison; web resource; hypothesis generation

## Introduction

Ligand-bound viral protein structures can provide a direct starting point for interpreting antiviral design opportunities. A bound ligand may reveal conserved interaction patterns, exposed atoms that can be discussed as potential modification sites, and target-side structural features that motivate additional review. In practice, however, these signals are rarely available in a single interface. Researchers often move among structure viewers, ligand dictionaries, contact calculations, solvent-exposure files, and custom scripts before reaching a coherent view of whether a particular structure is worth further design discussion. [CITATION NEEDED: RCSB/PDB] [CITATION NEEDED: structural bioinformatics tools]

Induced-proximity and degrader-oriented design questions add a further layer of interpretation. A viral target context may have a ligand-bound structure, but degrader-oriented follow-up also requires cautious review of ligand attachment opportunities, target-side exposed lysine cues, and the distinction between structural plausibility and experimental degradability. The useful question at this stage is not whether a system is proven degradable, but whether the available structural evidence supports a transparent, reviewable hypothesis for follow-up. [CITATION NEEDED: PROTAC review] [CITATION NEEDED: targeted protein degradation review]

V-LiSEMOD addresses this gap as a structure-guided viral protein-ligand exploration platform. It integrates curated viral protein metadata, ligand-centered interaction evidence, solvent-exposed atom review, functional-group annotations, cross-structure ligand comparison, and PROTACability-style heuristic triage. The manuscript-safe contribution is a practical evidence interface for hypothesis generation, design triage, and companion-tool context. Figure 1 situates V-LiSEMOD within a broader companion-tool ecosystem, while later figures focus on public workflow evidence and V-LiSEMOD-specific modules.

## Software Overview

V-LiSEMOD is organized around several user-facing workflows. The Structure Explorer supports virus, PDB, chain, and ligand selection, then provides ligand imagery, PyMOL-oriented session generation, binding-pocket context when available, SASA-derived exposed-atom interpretation, and functional-group visualization. Protein Query supports target-centric filtering and export-oriented dataset assembly. Ligand Indexer supports ligand-first lookup across mapped viral structures. Ligand Comparison lets users review how one ligand appears across multiple structural contexts. PROTACability Assessment provides transparent structural-priority and degrader-readiness triage across target, protein, structure, and chain-level views. An optional assistant route can be enabled in selected deployments, but it is treated as deployment-dependent and is not central to this manuscript.

The current public-facing workflow is documented as no-login by default, with public pages for the main scientific routes and a health endpoint captured as lightweight evidence. Companion-tool links support workflow continuation into related resources such as Warhead Hunter, PROTAC Builder, E3 Ligandalyzer, and PyMACS. These resources are manuscript context and continuation paths, not proof of a single unified backend. Figure 3 shows Warhead Hunter as a public exposure-oriented companion workflow, Figure 5 shows PROTAC Builder as downstream continuation context, and Figure 6 shows E3 Ligandalyzer and PyMACS as adjacent companion resources.

## Implementation and Architecture

The repository documentation describes V-LiSEMOD as a Flask-based web application. The main route areas include the home Structure Explorer, About, Protein Query, Ligand Indexer, Ligand Comparison, PROTACability Assessment, and `/healthz`. The application can operate against a local SQLite database, `viral_data.db`, or selected RANDY-backed route groups when the deployment is configured for that mode. Generated assets include charts, ligand images, coordinate caches, ligand SDF caches, PyMOL-oriented session files, and export outputs. These generated outputs are runtime artifacts rather than canonical manuscript data.

The database layer supports structure lookup, ligand identity, atom-level context, contact review, solvent exposure, functional-group annotation, and PROTACability-style triage. Representative tables include `Virus_Proteins`, `ligand_synonyms`, `ligand_atoms`, `Arpeggio_Contacts_Data`, `Ligand_Arp_Diagram`, `RUPLEY_SASA_DATA`, `Functional_GROUPED`, `Functional_Group_Atoms`, `protacability_assessment`, `protacability_lysine_proximity`, `protacability_ligand_inventory`, `protacability_warhead_linkability`, and `protacability_degrader_readiness`. These data layers support interface claims about what can be displayed or reviewed. They do not validate biological conclusions by themselves.

The optional Drug GPT/BioGPT route is controlled by feature flags and runtime dependencies. In the current screenshot package, `/drugapp/` returned a disabled or unavailable state and is therefore described only as optional deployment-dependent functionality. This conservative treatment keeps the manuscript centered on structure-guided viral ligand review rather than presenting optional assistant behavior as a core validated feature.

## User Workflows and Feature Demonstration

A structure-specific workflow begins with virus, PDB, and ligand selection. The user can inspect ligand context, generate ligand imagery, review exposed atoms when solvent-accessibility evidence exists, and export PyMOL-oriented material for deeper structural review. This workflow is designed to help users ask which ligand atoms appear available for discussion as modification sites while preserving important interactions.

A target-centric workflow begins with a protein or viral target question. Protein Query narrows relevant structures and links to export-ready data layers, giving reviewers a practical route from target name to PDB-linked ligand context. A ligand-first workflow begins with a ligand code or synonym, then maps the ligand into available structural contexts and interaction charts. A comparison workflow asks whether the same ligand retains or changes its interaction behavior across structures. Together these modules support design review without claiming global ligand intelligence or experimentally validated structure-activity relationships.

The PROTACability workflow is best interpreted as transparent heuristic triage. It combines ligand-centered warhead linkability, target-side lysine accessibility, structural-priority cues, simplified geometry context, and combined degrader-readiness summaries. These outputs prioritize structures for follow-up review. They should not be described as validated degrader prediction, guaranteed PROTAC design, or proof of productive ternary-complex formation. Figure 4 is the central V-LiSEMOD workflow figure and includes a PROTACability-style triage panel.

Companion-tool handoff gives users a way to continue from upstream structural review into adjacent design environments. Warhead Hunter contributes exposure-oriented public workflow context, PROTAC Builder provides a downstream assembly workspace, E3 Ligandalyzer provides recruiter-centered context, and PyMACS contributes a related computational resource. The appropriate manuscript framing is companion-tool context and workflow continuity, not full technical unification.

## Tables

{chr(10).join(table_sections)}

## Figures

{figure_sections}

## Validation and Reproducibility

The validation plan separates local run validation, data validation, workflow validation, manuscript evidence review, and reproducibility checklist items. Local validation should install dependencies from `requirements.txt` or `environment.yml`, start the app with `python app.py`, confirm `/healthz`, and verify the major public routes. If the optional assistant is discussed, both enabled and disabled behavior should be verified.

Data validation should confirm that the configured data source is present and that the expected core tables are available. In local SQLite mode this means checking `viral_data.db`; in RANDY-backed mode it means confirming the configured route groups respond as expected. PROTACability validation should confirm the assessment, lysine proximity, ligand inventory, warhead linkability, and degrader-readiness tables where those layers are described.

Workflow validation should exercise Structure Explorer, Protein Query, Ligand Indexer, Ligand Comparison, PROTACability Assessment, companion-tool handoff, and PyMOL-oriented export. The current screenshot package documents public screenshots captured on 2026-06-11 for Warhead Hunter, V-LiSEMOD, PROTAC Builder, E3 Ligandalyzer, PyMACS, and Warhead Hunter API documentation. Lightweight endpoint evidence includes Warhead Hunter health, manifest, and examples payloads plus V-LiSEMOD `/healthz`. These checks support reproducibility-oriented documentation, but they are not production-scale API guarantees.

## Limitations

V-LiSEMOD depends on available co-crystal structures, curated metadata, ligand mapping, solvent-exposure calculations, functional-group annotations, and generated runtime assets. The platform should therefore be described as curated and structure-guided rather than exhaustive. Absence of a strong heuristic score does not prove that a target is not degradable, and presence of a strong score does not guarantee successful degrader design.

PROTACability is a manuscript-sensitive term. In this draft it refers to transparent structural-priority and design-readiness heuristics. It does not mean experimentally validated targeted degradation, productive ternary-complex prediction, guaranteed warhead tractability, automated medicinal chemistry decision-making, or a substitute for biochemical and cellular assays.

Several implementation limitations should remain visible. Local or provisioned data assets are required for complete operation. Generated outputs depend on writable runtime directories. RANDY-backed operation is deployment-specific. The optional assistant route is not a default validated workflow. The AutoDock-Vina PrepServer figure slot remains pending because the current workspace is V-LiSEMOD rather than the expected local docking-preparation application. Some draft screenshots may need recapture because overlays, modals, or full-page capture artifacts reduce publication readiness.

## Availability and Data Access

Code availability: [REPOSITORY URL TO CONFIRM]

Web availability: [PUBLIC URLS TO CONFIRM]

Data availability: [DATA AVAILABILITY LANGUAGE TO CONFIRM]

License: [LICENSE TO CONFIRM]

Software dependencies and framework citations: [CITATION NEEDED: Flask/software framework] [CITATION NEEDED: SQLite] [CITATION NEEDED: PyMOL if cited] [CITATION NEEDED: Arpeggio/contact analysis if cited]

The public repository should not include private local databases, generated caches, credentials, tokens, local model weights, or unredacted runtime artifacts unless deliberately curated for release.

## Conclusions

V-LiSEMOD provides a practical, transparent platform for structure-guided viral ligand review. By bringing target metadata, ligand identity, interaction evidence, solvent-exposed atom context, functional-group annotations, ligand comparison, and heuristic PROTACability-style triage into a single workflow, it helps researchers move from ligand-bound viral structures to reviewable hypotheses and companion-tool handoff. The strongest manuscript framing is therefore not prediction or automation, but interpretability: V-LiSEMOD helps users organize evidence, prioritize follow-up review, and keep design claims separate from experimental validation.

## References

References are placeholders and require curator review before submission.

1. [CITATION NEEDED: RCSB/PDB or relevant structural database]
2. [CITATION NEEDED: Flask/software framework]
3. [CITATION NEEDED: SQLite or database layer if cited]
4. [CITATION NEEDED: PyMOL if export workflow is discussed]
5. [CITATION NEEDED: Arpeggio or contact-analysis method if cited]
6. [CITATION NEEDED: solvent-accessible surface area method]
7. [CITATION NEEDED: PROTAC review]
8. [CITATION NEEDED: targeted protein degradation review]
9. [CITATION NEEDED: viral structural bioinformatics or antiviral ligand-design context]

## Figure Captions

{chr(10).join(f'**Figure {fig.number}. {fig.title}.** {fig.caption}' for fig in FIGURES)}

## Table Captions

1. Platform capabilities: Current V-LiSEMOD modules and their primary user-facing roles.
2. Data-layer summary: Major V-LiSEMOD data layers supporting structure, ligand, interaction, solvent-exposure, and PROTACability workflows.
3. PROTACability interpretation guardrails: Interpretation boundaries for heuristic evidence layers.
4. Module-to-user-question map: Representative research questions supported by each V-LiSEMOD workflow.
5. Limitations and future work: Current limitations and future documentation or validation directions.
6. Reproducibility checklist: Items needed for reproducible local or provisioned V-LiSEMOD runs.

## Reviewer and Collaborator Notes

- The draft uses manuscript-safe language: heuristic triage, hypothesis generation, structure-guided review, representative screenshots, and companion-tool context.
- The draft intentionally avoids claims of validated degrader prediction, guaranteed PROTAC design, exhaustive viral ligand coverage, full technical unification, and production-scale API guarantees.
- Bracketed citation placeholders require curator review and replacement with real references.
- Figure panels are representative screenshots and draft crops, not final journal-quality art.
- The AutoDock-Vina PrepServer figure slot is intentionally left as a placeholder until the correct application context is available.
- Public endpoint checks are evidence of reachability at capture time, not long-term availability commitments.
"""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for margin_name, width in (("top", "80"), ("left", "120"), ("bottom", "80"), ("right", "120")):
        node = OxmlElement(f"w:{margin_name}")
        node.set(qn("w:w"), width)
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(text)


def add_table(doc, title: str, rows: list[list[str]], number: int) -> None:
    add_caption(doc, f"Table {number}. {title}.")
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    set_cell_margins(table)
    widths = [1.65, 2.25, 2.6] if len(rows[0]) == 3 else [6.5 / len(rows[0])] * len(rows[0])
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        hdr[idx].text = value
        set_cell_shading(hdr[idx], "F2F4F7")
        for paragraph in hdr[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_values in rows[1:]:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def image_width_inches(path: Path, max_width: float = 6.15) -> float:
    if Image is None:
        return max_width
    with Image.open(path) as img:
        w, h = img.size
    if h == 0 or w == 0:
        return max_width
    if w > h * 1.2:
        return max_width
    return min(max_width, 4.8)


def add_figure(doc, fig: Figure) -> None:
    doc.add_heading(f"Figure {fig.number}. {fig.title}", level=3)
    if fig.placeholder:
        p = doc.add_paragraph()
        p.add_run("Draft figure placeholder: ").bold = True
        p.add_run(fig.placeholder)
    inserted = False
    for file_name in fig.files:
        path = FIG_DIR / file_name
        if path.exists():
            doc.add_picture(str(path), width=Inches(image_width_inches(path)))
            inserted = True
        else:
            p = doc.add_paragraph()
            p.add_run("Missing image: ").bold = True
            p.add_run(f"docs/manuscript/figures/{file_name}")
    if not inserted and not fig.placeholder:
        doc.add_paragraph("[Draft figure placeholder: image pending.]")
    add_caption(doc, f"Figure {fig.number}. {fig.title}. {fig.caption}")


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("V-LiSEMOD: structure-guided viral ligand analysis with transparent degrader-readiness triage")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    for text in [
        "Draft manuscript for collaborator review",
        "Authors: [AUTHOR NAMES TO ADD]",
        "Affiliations: [AFFILIATIONS TO ADD]",
        "Corresponding author: [CONTACT TO ADD]",
        "Running title: V-LiSEMOD viral ligand analysis platform",
        "Article type: software article, webserver paper, application note, methods article, or resource article",
        "Reference status: citations are placeholders and require curator review.",
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)


def build_docx() -> None:
    if Document is None:
        raise RuntimeError(f"python-docx is unavailable: {IMPORT_ERROR}")

    doc = Document()
    configure_styles(doc)
    add_title_page(doc)

    doc.add_heading("Abstract", level=1)
    add_paragraphs(
        doc,
        [
            "Viral protein-ligand co-crystal structures contain practical signals for antiviral chemical biology and medicinal chemistry, but design-relevant evidence is often distributed across structural files, interaction tables, solvent accessibility calculations, ligand identifiers, and downstream scripts. V-LiSEMOD is a structure-guided viral protein-ligand exploration platform that brings curated viral target metadata, ligand-centered interaction evidence, solvent-exposed atom review, functional-group annotations, cross-structure ligand comparison, and PROTACability-style heuristic triage into a web-accessible workflow. The platform supports structure-specific exploration, target-centric query, ligand-first lookup, ligand comparison, PyMOL-oriented export, and companion-tool handoff for downstream degrader-oriented design exploration. Its PROTACability layer is framed as transparent hypothesis generation: warhead linkability, target-side lysine accessibility, structural-priority cues, and combined degrader-readiness summaries help users prioritize contexts for review without claiming experimental degradation prediction. Representative screenshots and lightweight endpoint evidence document the current public-facing workflow, while validation notes outline reproducibility checks for local SQLite-backed and optional RANDY-backed deployments. V-LiSEMOD is intended to help researchers move from viral co-crystal context to interpretable design questions, not to replace biochemical validation, cellular degradation assays, ternary-complex modeling, or medicinal chemistry judgment.",
        ],
    )
    doc.add_paragraph("Keywords: viral ligand analysis; structural bioinformatics; solvent exposure; PROTACability; degrader design; ligand comparison; web resource; hypothesis generation")

    sections = [
        ("Introduction", [
            "Ligand-bound viral protein structures can provide a direct starting point for interpreting antiviral design opportunities. A bound ligand may reveal conserved interaction patterns, exposed atoms that can be discussed as potential modification sites, and target-side structural features that motivate additional review. [CITATION NEEDED: RCSB/PDB]",
            "Induced-proximity and degrader-oriented design questions add a further layer of interpretation. A viral target context may have a ligand-bound structure, but degrader-oriented follow-up also requires cautious review of ligand attachment opportunities, target-side exposed lysine cues, and the distinction between structural plausibility and experimental degradability. [CITATION NEEDED: PROTAC review]",
            "V-LiSEMOD addresses this gap as a structure-guided viral protein-ligand exploration platform for evidence review, heuristic triage, and companion-tool handoff.",
        ]),
        ("Software Overview", [
            "V-LiSEMOD is organized around Structure Explorer, Protein Query, Ligand Indexer, Ligand Comparison, PROTACability Assessment, companion-tool handoff, and an optional deployment-dependent assistant route.",
            "The main scientific workflows are documented as public and no-login by default in the current deployment notes. Companion tools provide context and downstream continuation, not evidence of a single unified backend.",
        ]),
        ("Implementation and Architecture", [
            "The repository documentation describes V-LiSEMOD as a Flask-based web application with local SQLite-backed operation and optional RANDY-backed route groups. Generated assets include charts, ligand images, coordinate caches, ligand SDF caches, PyMOL-oriented sessions, and export outputs.",
            "The database layer supports curated target metadata, ligand identity, atom-level context, interaction evidence, solvent-exposure cues, functional-group annotation, and PROTACability-style triage. These layers support review workflows but do not validate biological conclusions by themselves.",
        ]),
        ("User Workflows and Feature Demonstration", [
            "A structure-specific workflow begins with virus, PDB, and ligand selection, then moves into ligand imagery, exposed-atom context, functional-group annotations, and PyMOL-oriented review. Target-centric, ligand-first, and comparison workflows support complementary entry points for review.",
            "The PROTACability workflow combines ligand-centered and target-centered evidence layers as transparent heuristic triage. The appropriate interpretation is prioritization for follow-up review, not experimental degradation prediction.",
        ]),
    ]
    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        add_paragraphs(doc, paragraphs)

    doc.add_heading("Tables", level=1)
    for number, (title, rows) in enumerate(TABLES.items(), start=1):
        add_table(doc, title, rows, number)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Figures", level=1)
    for fig in FIGURES:
        add_figure(doc, fig)

    trailing_sections = [
        ("Validation and Reproducibility", [
            "The validation plan separates local run validation, data validation, workflow validation, manuscript evidence review, and reproducibility checklist items. Local validation should confirm `/healthz` and the major public routes.",
            "The current screenshot package documents public screenshots captured on 2026-06-11 and lightweight endpoint evidence for Warhead Hunter and V-LiSEMOD. These checks are reachability and reproducibility aids, not production-scale API guarantees.",
        ]),
        ("Limitations", [
            "V-LiSEMOD depends on available co-crystal structures, curated metadata, ligand mapping, solvent-exposure calculations, functional-group annotations, and generated runtime assets. It should not be described as exhaustive.",
            "PROTACability refers to transparent structural-priority and design-readiness heuristics. It does not mean experimentally validated targeted degradation, productive ternary-complex prediction, guaranteed warhead tractability, or automated medicinal chemistry decision-making.",
            "The AutoDock-Vina PrepServer figure slot remains pending because the current workspace is V-LiSEMOD rather than the expected local docking-preparation application.",
        ]),
        ("Availability and Data Access", [
            "Code availability: [REPOSITORY URL TO CONFIRM]",
            "Web availability: [PUBLIC URLS TO CONFIRM]",
            "Data availability: [DATA AVAILABILITY LANGUAGE TO CONFIRM]",
            "License: [LICENSE TO CONFIRM]",
            "Software dependencies and framework citations: [CITATION NEEDED: Flask/software framework] [CITATION NEEDED: SQLite] [CITATION NEEDED: PyMOL if cited] [CITATION NEEDED: Arpeggio/contact analysis if cited]",
        ]),
        ("Conclusions", [
            "V-LiSEMOD provides a practical, transparent platform for structure-guided viral ligand review. Its contribution is evidence organization, hypothesis generation, design triage, and companion-tool context rather than prediction or automation.",
        ]),
        ("References", [
            "References are placeholders and require curator review before submission.",
            "1. [CITATION NEEDED: RCSB/PDB or relevant structural database]",
            "2. [CITATION NEEDED: Flask/software framework]",
            "3. [CITATION NEEDED: SQLite or database layer if cited]",
            "4. [CITATION NEEDED: PyMOL if export workflow is discussed]",
            "5. [CITATION NEEDED: Arpeggio or contact-analysis method if cited]",
            "6. [CITATION NEEDED: solvent-accessible surface area method]",
            "7. [CITATION NEEDED: PROTAC review]",
            "8. [CITATION NEEDED: targeted protein degradation review]",
            "9. [CITATION NEEDED: viral structural bioinformatics or antiviral ligand-design context]",
        ]),
        ("Figure Captions", [f"Figure {fig.number}. {fig.title}. {fig.caption}" for fig in FIGURES]),
        ("Table Captions", [
            "Table 1. Platform capabilities: Current V-LiSEMOD modules and their primary user-facing roles.",
            "Table 2. Data-layer summary: Major V-LiSEMOD data layers supporting structure, ligand, interaction, solvent-exposure, and PROTACability workflows.",
            "Table 3. PROTACability interpretation guardrails: Interpretation boundaries for heuristic evidence layers.",
            "Table 4. Module-to-user-question map: Representative research questions supported by each V-LiSEMOD workflow.",
            "Table 5. Limitations and future work: Current limitations and future documentation or validation directions.",
            "Table 6. Reproducibility checklist: Items needed for reproducible local or provisioned V-LiSEMOD runs.",
        ]),
        ("Reviewer and Collaborator Notes", [
            "This draft intentionally uses cautious language: heuristic triage, hypothesis generation, structure-guided review, representative screenshots, and companion-tool context.",
            "Citation placeholders, figure quality, public URLs, repository URL, license, data availability wording, and author/affiliation metadata need manual review.",
        ]),
    ]
    for heading, paragraphs in trailing_sections:
        doc.add_heading(heading, level=1)
        add_paragraphs(doc, paragraphs)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def build_notes() -> str:
    inserted = []
    missing = []
    for fig in FIGURES:
        if fig.placeholder:
            missing.append(f"Figure {fig.number}: {fig.placeholder}")
        for file_name in fig.files:
            if (FIG_DIR / file_name).exists():
                inserted.append(f"Figure {fig.number}: docs/manuscript/figures/{file_name}")
            else:
                missing.append(f"Figure {fig.number}: docs/manuscript/figures/{file_name}")

    citation_placeholders = [
        "RCSB/PDB or relevant structural database",
        "Flask/software framework",
        "SQLite or database layer",
        "PyMOL if export workflow is discussed",
        "Arpeggio/contact-analysis method if cited",
        "solvent-accessible surface area method",
        "PROTAC review",
        "targeted protein degradation review",
        "viral structural bioinformatics or antiviral ligand-design context",
    ]

    return f"""# Manuscript Draft Notes

## Source Files Used

{chr(10).join(f'- `{path}`' for path in SOURCE_FILES if (ROOT / path).exists())}

## Figures Inserted

{chr(10).join(f'- {item}' for item in inserted)}

## Figures Omitted or Missing

{chr(10).join(f'- {item}' for item in missing)}

## Tables Included

{chr(10).join(f'- {title}' for title in TABLES)}

## Claims Intentionally Avoided

- Experimentally validated degradation prediction
- Guaranteed PROTAC design success
- Productive ternary-complex prediction
- Automated medicinal chemistry decision-making
- Exhaustive viral ligand or target coverage
- Full technical unification across companion tools
- Production-scale API guarantees or service-level commitments

## Citation Placeholders Needing Follow-up

{chr(10).join(f'- [CITATION NEEDED: {item}]' for item in citation_placeholders)}

## Validation Notes

- The DOCX should be validated after regeneration with `unzip -t docs/manuscript/drafts/V-LiSEMOD_manuscript_draft.docx`.
- The DOCX should be loaded with `python-docx` to confirm paragraphs, tables, and inline images are present.
- A render attempt with the bundled `render_docx.py` may fail on this workstation if headless LibreOffice cannot load `/opt/homebrew/opt/little-cms2/lib/liblcms2.2.dylib`; that is a local renderer dependency blocker, not evidence that the DOCX package is corrupt.
- If LibreOffice is repaired or available elsewhere, rerun DOCX-to-PNG rendering and visually inspect all pages before journal submission.

## Manual Review Checklist

- Add real author names, affiliations, contact details, funding, acknowledgments, conflicts, and ethics statements as appropriate.
- Confirm final repository URL, public URLs, license, and data availability language.
- Replace citation placeholders with verified references.
- Recheck all figure panels for publication readiness, overlays, modals, tiny text, and private UI state.
- Regenerate the pending AutoDock-Vina PrepServer panel from the correct application context.
- Re-run local app validation and endpoint checks against the final manuscript version.
- Confirm PROTACability wording remains explicitly heuristic throughout the manuscript.
- Decide whether the article target is a software article, webserver paper, application note, methods article, or resource article.
"""


def main() -> None:
    DRAFT_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(build_markdown(), encoding="utf-8")
    NOTES_PATH.write_text(build_notes(), encoding="utf-8")
    build_docx()
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {NOTES_PATH}")


if __name__ == "__main__":
    main()
